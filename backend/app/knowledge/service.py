import hashlib
import math
import re
import json

from sqlalchemy.orm import Session

from app.knowledge.embeddings import embed_texts
from app.models import KnowledgeChunk, KnowledgeDoc

EMBED_DIM = 64
DEFAULT_MAX_CHARS = 600
OVERLAP_CHARS = 80
SUPPORTED_TEXT_TYPES = {"txt", "md", "markdown", "csv", "json", "pdf", "docx", "xlsx", "xls"}


def _tokenize(text: str) -> list[str]:
    tokens: list[str] = []
    for match in re.findall(r"[A-Za-z0-9_]+|[\u4e00-\u9fff]+", text.lower()):
        if re.fullmatch(r"[\u4e00-\u9fff]+", match):
            tokens.extend(list(match))
            if len(match) >= 2:
                tokens.extend(match[i : i + 2] for i in range(len(match) - 1))
        else:
            tokens.append(match)
    return tokens


def local_embedding(text: str) -> list[float]:
    vector = [0.0] * EMBED_DIM
    for token in _tokenize(text):
        digest = int(hashlib.md5(token.encode("utf-8")).hexdigest()[:8], 16)
        vector[digest % EMBED_DIM] += 1.0
    norm = math.sqrt(sum(v * v for v in vector)) or 1.0
    return [v / norm for v in vector]


def cosine_similarity(a: list[float], b: list[float]) -> float:
    return sum(x * y for x, y in zip(a, b))


def chunk_text(text: str, max_chars: int = DEFAULT_MAX_CHARS, overlap: int = OVERLAP_CHARS) -> list[str]:
    normalized = re.sub(r"\s+", " ", text).strip()
    if not normalized:
        return []
    chunks: list[str] = []
    start = 0
    while start < len(normalized):
        end = min(start + max_chars, len(normalized))
        chunks.append(normalized[start:end])
        if end == len(normalized):
            break
        start = max(start + max_chars - overlap, start + 1)
    return chunks


def _parse_pdf(data: bytes) -> str:
    """Extract text from PDF using pypdf."""
    import io
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text.strip())
    return "\n\n".join(parts)


def _parse_docx(data: bytes) -> str:
    """Extract text from Word .docx using python-docx."""
    import io
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _parse_xlsx(data: bytes) -> str:
    """Extract data from Excel .xlsx/.xls using openpyxl."""
    import io
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        parts.append(f"[Sheet: {ws.title}]")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append("\t".join(cells))
    wb.close()
    return "\n".join(parts)


def parse_text_file(name: str, data: bytes) -> str:
    suffix = name.rsplit(".", 1)[-1].lower() if "." in name else ""
    if suffix not in SUPPORTED_TEXT_TYPES:
        raise ValueError(f"暂不支持的文件类型: .{suffix}")
    if suffix in {"txt", "md", "markdown", "csv", "json"}:
        return data.decode("utf-8", errors="replace")
    if suffix == "pdf":
        return _parse_pdf(data)
    if suffix == "docx":
        return _parse_docx(data)
    if suffix in {"xlsx", "xls"}:
        return _parse_xlsx(data)
    return data.decode("utf-8", errors="replace")


class KnowledgeService:
    def __init__(self, db: Session) -> None:
        self.db = db

    def ingest(
        self,
        tenant_id: str,
        name: str,
        content: str,
        file_type: str,
        size_bytes: int,
        industry_id: str | None = None,
    ) -> KnowledgeDoc:
        doc = KnowledgeDoc(
            tenant_id=tenant_id,
            industry_id=industry_id,
            name=name,
            file_type=file_type,
            size_bytes=size_bytes,
            status="processing",
        )
        self.db.add(doc)
        self.db.flush()

        chunks = chunk_text(content)
        embeddings = embed_texts(chunks)
        for index, chunk in enumerate(chunks):
            embedding = embeddings[index] if index < len(embeddings) else local_embedding(chunk)
            self.db.add(
                KnowledgeChunk(
                    tenant_id=tenant_id,
                    industry_id=industry_id,
                    doc_id=doc.id,
                    chunk_index=index,
                    content=chunk,
                    embedding_json=embedding,
                    metadata_json={"doc_name": name, "chunk_index": index},
                    tokens=len(_tokenize(chunk)),
                )
            )
        doc.status = "ready"
        doc.chunk_count = len(chunks)
        self.db.commit()
        self.db.refresh(doc)
        return doc

    def search(self, tenant_id: str, query: str, top_k: int = 5, industry_id: str | None = None) -> list[dict]:
        query_filter = [KnowledgeChunk.tenant_id == tenant_id]
        if industry_id:
            query_filter.append(KnowledgeChunk.industry_id == industry_id)
        chunks = (
            self.db.query(KnowledgeChunk)
            .filter(*query_filter)
            .order_by(KnowledgeChunk.created_at.desc())
            .all()
        )
        query_vec = embed_texts([query])[0]
        query_tokens = set(_tokenize(query))
        by_doc: dict[str, list[tuple[float, KnowledgeChunk]]] = {}
        for chunk in chunks:
            vec = chunk.embedding_json or local_embedding(chunk.content)
            vector_score = cosine_similarity(query_vec, vec)
            overlap = len(query_tokens & set(_tokenize(chunk.content)))
            score = vector_score * 0.7 + min(overlap, 5) * 0.06
            by_doc.setdefault(chunk.doc_id, []).append((score, chunk))

        ranked: list[tuple[float, list[tuple[float, KnowledgeChunk]]]] = []
        for doc_id, items in by_doc.items():
            best_score = max(score for score, _ in items)
            # 超长表格降权，避免重复行淹没精炼知识
            if len(items) > 100:
                size_penalty = 0.45
            elif len(items) > 30:
                size_penalty = 0.7
            else:
                size_penalty = 1.0
            metadata = items[0][1].metadata_json or {}
            doc_name = metadata.get("doc_name", "")
            name_overlap = len(query_tokens & set(_tokenize(doc_name)))
            effective = best_score * size_penalty + name_overlap * 0.05
            ranked.append((effective, items))
        ranked.sort(key=lambda item: item[0], reverse=True)

        results: list[dict] = []
        for effective_score, items in ranked[:top_k]:
            items.sort(key=lambda item: item[0], reverse=True)
            score, chunk = items[0]
            results.append(
                {
                    "id": chunk.id,
                    "doc_id": chunk.doc_id,
                    "content": chunk.content,
                    "score": round(effective_score, 4),
                    "metadata": chunk.metadata_json,
                }
            )
        return results

    def batch_search(
        self,
        tenant_id: str,
        queries: list[tuple[str, int]],
        industry_id: str | None = None,
    ) -> list[list[dict]]:
        """Batch search: load chunks once, embed all queries in one call, return per-query results."""
        if not queries:
            return []
        query_texts = [q[0] for q in queries]
        top_ks = [q[1] for q in queries]

        # Load chunks once via raw SQL (much faster than ORM full-object load)
        from sqlalchemy import text as sql_text
        if industry_id:
            rows = self.db.execute(
                sql_text("SELECT id, doc_id, content, embedding_json, metadata_json FROM knowledge_chunks WHERE tenant_id = :tid AND industry_id = :iid"),
                {"tid": tenant_id, "iid": industry_id},
            ).fetchall()
        else:
            rows = self.db.execute(
                sql_text("SELECT id, doc_id, content, embedding_json, metadata_json FROM knowledge_chunks WHERE tenant_id = :tid"),
                {"tid": tenant_id},
            ).fetchall()
        if not rows:
            return [[] for _ in queries]

        # Batch-embed all queries in one Ollama call
        query_vecs = embed_texts(query_texts)

        # Precompute chunk vectors and token sets once
        chunk_data = []
        for row in rows:
            emb_str = row[3]
            if emb_str:
                vec = json.loads(emb_str) if isinstance(emb_str, str) else emb_str
            else:
                vec = local_embedding(row[2])
            chunk_data.append({
                "id": row[0],
                "doc_id": row[1],
                "content": row[2],
                "vec": vec,
                "tokens": set(_tokenize(row[2])),
                "metadata": json.loads(row[4]) if row[4] and isinstance(row[4], str) else (row[4] or {}),
            })

        # Group by doc for penalty
        doc_groups: dict[str, list] = {}
        for cd in chunk_data:
            doc_groups.setdefault(cd["doc_id"], []).append(cd)

        all_results = []
        for qi, (qvec, top_k) in enumerate(zip(query_vecs, top_ks)):
            q_tokens = set(_tokenize(query_texts[qi]))
            ranked = []
            for doc_id, items in doc_groups.items():
                best_score = -1.0
                for cd in items:
                    vs = cosine_similarity(qvec, cd["vec"])
                    overlap = len(q_tokens & cd["tokens"])
                    score = vs * 0.7 + min(overlap, 5) * 0.06
                    if score > best_score:
                        best_score = score
                if len(items) > 100:
                    size_penalty = 0.45
                elif len(items) > 30:
                    size_penalty = 0.7
                else:
                    size_penalty = 1.0
                metadata = items[0]["metadata"]
                doc_name = metadata.get("doc_name", "")
                name_overlap = len(q_tokens & set(_tokenize(doc_name)))
                effective = best_score * size_penalty + name_overlap * 0.05
                ranked.append((effective, items))
            ranked.sort(key=lambda x: x[0], reverse=True)
            results = []
            for effective_score, items in ranked[:top_k]:
                items_sorted = sorted(items, key=lambda cd: cosine_similarity(qvec, cd["vec"]), reverse=True)
                best = items_sorted[0]
                results.append({
                    "id": best["id"],
                    "doc_id": best["doc_id"],
                    "content": best["content"],
                    "score": round(effective_score, 4),
                    "metadata": best["metadata"],
                })
            all_results.append(results)
        return all_results

    def list_docs(self, tenant_id: str) -> list[dict]:
        rows = (
            self.db.query(KnowledgeDoc)
            .filter(KnowledgeDoc.tenant_id == tenant_id)
            .order_by(KnowledgeDoc.created_at.desc())
            .all()
        )
        return [
            {
                "id": row.id,
                "name": row.name,
                "file_type": row.file_type,
                "status": row.status,
                "chunk_count": row.chunk_count,
                "size_bytes": row.size_bytes,
                "created_at": row.created_at.isoformat() if row.created_at else None,
            }
            for row in rows
        ]

    def delete(self, tenant_id: str, doc_id: str) -> bool:
        doc = self.db.query(KnowledgeDoc).filter(
            KnowledgeDoc.id == doc_id, KnowledgeDoc.tenant_id == tenant_id
        ).first()
        if doc is None:
            return False
        self.db.query(KnowledgeChunk).filter(
            KnowledgeChunk.doc_id == doc_id, KnowledgeChunk.tenant_id == tenant_id
        ).delete(synchronize_session=False)
        self.db.delete(doc)
        self.db.commit()
        return True
